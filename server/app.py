import os
from flask import Flask, jsonify, request, redirect, url_for, make_response, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_bcrypt import Bcrypt
from flask_jwt_extended import JWTManager, create_access_token, get_jwt_identity, jwt_required
from flask_restful import Api, Resource
from datetime import timedelta
from sqlalchemy import func
from datetime import datetime, timezone
from werkzeug.utils import secure_filename
import spacy
import PyPDF2
import docx
from flask_cors import CORS
from docx import Document as DocxDocument
import io


from models import db, User, Document, Suggestion, DocumentHistory, StopWord, Word

# Initialize the Flask app
app = Flask(__name__)

CORS(app) 


# Configure the app for SQLite
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///refinedocs.db'  # Using SQLite
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = 'your_secret_key'  # Change to your secret key
app.config['JWT_SECRET_KEY'] = 'your_jwt_secret_key'  # Change to your JWT secret key
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(days=1)  # Token expiration (1 day)

# Initialize extensions

db.init_app(app) 
migrate = Migrate(app, db)
bcrypt = Bcrypt(app)
jwt = JWTManager(app)
api = Api(app)  # Initialize the Api for flask_restful

# Simple route to check if the server is running
@app.route('/')
def index():
    return jsonify({"message": "RefineDocs API is running!"})

# CHECKSESSION

class CheckSession(Resource):
    @jwt_required()
    def get(self):
        user_id = get_jwt_identity()
        if user_id:
            user = User.query.filter_by(id=user_id).first()
            if user:
                user_data = user.to_dict()
                print("User Data:", user_data)  # Print the user data
                return {'user_type': 'user', 'user_data': user_data}, 200
        
        return jsonify({'message': '401: Not Authorized'}), 401

api.add_resource(CheckSession, '/check_session')


# ROUTES FOR USER

@app.route('/users/registration', methods=['POST'])
def register_user():
    if request.method == 'POST':
        data = request.get_json()
        username = data.get('username')
        email = data.get('email').strip().lower()
        password = data.get('password')

        # Check if the email or username is already registered
        existing_user = User.query.filter(
            (func.lower(User.email) == email) | (User.username == username)
        ).first()
        
        if existing_user:
            error_message = "Email or Username already registered. Please use a different one."
            return jsonify({'error': error_message}), 400
        
        # Create new user
        new_user = User(
            username=username, 
            email=email, 
            password_hash=bcrypt.generate_password_hash(password).decode('utf-8')
        )

        db.session.add(new_user)
        db.session.commit()

        access_token = create_access_token(identity=new_user.id, expires_delta=timedelta(hours=24))
        return jsonify({'access_token': access_token}), 201

    return redirect(url_for('index'))


@app.route('/users/signin', methods=['POST'])
def signin_user():
    if request.method == 'POST':
        data = request.get_json()
        email = data.get('email').strip().lower()
        password = data.get('password')

        user = User.query.filter(func.lower(User.email) == email).first()

        if user and bcrypt.check_password_hash(user.password_hash, password):
            access_token = create_access_token(identity=user.id, expires_delta=timedelta(hours=24))
            return jsonify({'access_token': access_token}), 200
        else:
            error_message = "Invalid email or password."
            return jsonify({'error': error_message}), 401

    return redirect(url_for('index'))


@app.route('/users/logout', methods=['DELETE'])
def logout_user():
    # JWT does not store sessions server-side, so there is no need to manage server-side session here.
    # Simply inform the client to discard the token.
    logout_dict = {'message': 'Logout successful. Goodbye!'}
    response = make_response(
        jsonify(logout_dict),
        200
    )
    return response


# EXTRACT FILE

def extract_text(filepath):
    _, file_extension = os.path.splitext(filepath)

    if file_extension.lower() == '.txt':
        with open(filepath, 'r') as file:
            return file.read()
    elif file_extension.lower() == '.pdf':
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ''
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text()
            return text
    elif file_extension.lower() == '.docx':
        doc = docx.Document(filepath)
        return '\n'.join([para.text for para in doc.paragraphs])
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


# IMPROVE DOCUMENTS.

def improve_document(original_content):
    # Load the spaCy NLP model
    nlp = spacy.load("en_core_web_sm")
    
    # Process the text through the NLP pipeline
    doc = nlp(original_content)
    
    # Improved content (initially the same as original content)
    improved_content = original_content
    
    # Example of extracting suggestions
    suggestions = []
    for token in doc:
        if token.is_stop:
            suggestions.append({
                "suggestion_text": f"Consider removing stopword: {token.text}",
                "original_word": token.text,  # Track the word being targeted
                "suggestion_type": "remove_stopword"
            })
        elif token.is_alpha and len(token.text) > 6:  # Example of a rule to improve word usage
            suggestions.append({
                "suggestion_text": f"Try simplifying: {token.text}",
                "original_word": token.text,  # Track the complex word
                "suggested_word": token.lemma_,  # Example: use lemma or another simpler word
                "suggestion_type": "simplify_word"
            })

    return improved_content, suggestions



# POST DOCUMENT

UPLOAD_FOLDER = 'assets'
ALLOWED_EXTENSIONS = {'txt', 'docx', 'pdf'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/documents/upload', methods=['POST'])
@jwt_required()
def upload_document():
    user_id = get_jwt_identity()
    user = User.query.get(user_id)

    if not user:
        return jsonify({'error': 'Unauthorized'}), 401

    if 'file' not in request.files:
        return jsonify({'error': 'No file selected'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Extract text from the uploaded document (based on the file type)
        original_content = extract_text(filepath)

        # Process the document with spaCy to improve it
        improved_content, suggestions = improve_document(original_content)

        # Create a new Document object and save it in the database
        new_document = Document(
            user_id=user.id,
            upload_date=datetime.now(timezone.utc),
            original_content=original_content,
            improved_content=improved_content
        )

        db.session.add(new_document)
        db.session.commit()

        # Add each suggestion to the suggestions table
        for suggestion_data in suggestions:
            new_suggestion = Suggestion(
                document_id=new_document.id,
                suggestion_text=suggestion_data['suggestion_text'],  # Only the suggestion text
                original_word=suggestion_data.get('original_word', None),  # Store the original word
                suggested_word=suggestion_data.get('suggested_word', None),  # Store the suggested word
                suggestion_type=suggestion_data.get('suggestion_type', None),  # Store suggestion type
                is_accepted=False  # Default to not accepted
            )
            db.session.add(new_suggestion)

        db.session.commit()

        return jsonify({
            'message': 'Document uploaded and processed successfully',
            'original_content': original_content,
            'improved_content': improved_content,
            'suggestions': suggestions
        }), 201

    return jsonify({'error': 'Invalid file format'}), 400


# FETCH DOCUMENT

@app.route('/documents/latest', methods=['GET'])
@jwt_required()  # Ensure the user is authenticated
def get_latest_document():
    user_id = get_jwt_identity()  # Get the current user's ID from the JWT token
    
    # Fetch the most recent document for the user
    latest_document = Document.query.filter_by(user_id=user_id).order_by(Document.upload_date.desc()).first()
    
    if not latest_document:
        return jsonify({'error': 'No documents found for this user'}), 404

    # Serialize the document data
    document_data = {
        'id': latest_document.id, 
        'original_content': latest_document.original_content,
        'improved_content': latest_document.improved_content,
        # 'upload_date': latest_document.upload_date.isoformat()  
    }

    return jsonify({'latest_document': document_data}), 200


# FETCH SUGGESTION

@app.route('/documents/<int:document_id>/suggestions', methods=['GET'])
@jwt_required()
def get_document_suggestions(document_id):
    # Check if the document exists
    document = Document.query.get(document_id)

    if not document:
        return jsonify({'error': 'Document not found'}), 404

    # Fetch all suggestions for the given document
    suggestions = Suggestion.query.filter_by(document_id=document_id).all()

    if not suggestions:
        return jsonify({'message': 'No suggestions found for this document'}), 404

    # Manually serialize suggestions using to_dict method
    suggestions_data = [suggestion.to_dict() for suggestion in suggestions]

    return jsonify({
        'document_id': document_id,
        'suggestions': suggestions_data
    }), 200


# ACCEPT SUGGESTION

@app.route('/suggestions/accept/<int:suggestion_id>', methods=['POST'])
@jwt_required()
def accept_suggestion(suggestion_id):
    user_id = get_jwt_identity()
    user = User.query.get(user_id)

    if not user:
        return jsonify({'error': 'Unauthorized'}), 401

    # Fetch the suggestion
    suggestion = Suggestion.query.get(suggestion_id)
    if not suggestion:
        return jsonify({'error': 'Suggestion not found'}), 404

    # Fetch the related document
    document = Document.query.get(suggestion.document_id)
    if not document:
        return jsonify({'error': 'Document not found'}), 404

    improved_content = document.improved_content

    # Handle suggestion based on its type
    if suggestion.suggestion_type == 'simplify_word':
        # Replace the original word with the suggested word in the improved content
        if suggestion.original_word and suggestion.suggested_word:
            improved_content = improved_content.replace(suggestion.original_word, suggestion.suggested_word)
    elif suggestion.suggestion_type == 'remove_stopword':
        # Remove the stopword from the improved content
        if suggestion.original_word:
            improved_content = improved_content.replace(suggestion.original_word, '')

    # Update the document's improved content and mark the suggestion as accepted
    document.improved_content = improved_content.strip()  # Strip to remove any extra spaces from removing stopwords
    suggestion.is_accepted = True

    try:
        db.session.commit()
        return jsonify({
            'message': 'Suggestion accepted and applied successfully',
            'updated_improved_content': improved_content
        }), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to accept suggestion: {str(e)}'}), 500


# DENY SUGGESTION

@app.route('/suggestions/deny/<int:suggestion_id>', methods=['DELETE'])
@jwt_required()
def deny_suggestion(suggestion_id):
    # Get the suggestion
    suggestion = Suggestion.query.get(suggestion_id)

    if not suggestion:
        return jsonify({'error': 'Suggestion not found'}), 404

    # Delete the suggestion
    db.session.delete(suggestion)
    db.session.commit()

    return jsonify({'message': 'Suggestion denied and deleted'}), 200


# EXPORT DOCUMENT


@app.route('/export-document/<int:document_id>', methods=['GET'])
@jwt_required()
def export_document(document_id):
    user_id = get_jwt_identity()  # Get the current user's ID from the JWT token
    
    # Fetch the document by ID and user
    document = Document.query.filter_by(id=document_id, user_id=user_id).first()

    if not document:
        return jsonify({'error': 'Document not found'}), 404

    # Create a new Word document using python-docx
    doc = DocxDocument()

    # Add content to the document (you can customize this)
    doc.add_heading('Improved Document', 0)
    doc.add_paragraph(document.improved_content)

    # Save the document to a BytesIO stream
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)

    # Return the Word document as a downloadable file
    return send_file(
        output_stream,
        as_attachment=True,
        download_name="improved_document.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# Run the app
if __name__ == '__main__':
    app.run(port=5555)
